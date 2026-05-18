from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT_DIR = Path(r"D:\Grade Three Down\ZhuanZ\campus_circle2.0\artifacts")
OUT_PATH = OUT_DIR / "project_midterm_reflection_fixed.docx"


def set_east_asia_font(run, name, size=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_east_asia_font(run, "黑体", 14 if level == 1 else 12, True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_east_asia_font(run, "宋体", 12)
    return p


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("项目中期心得总结")
    set_east_asia_font(run, "黑体", 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("以“校圈”校园信息与社区平台开发为例")
    set_east_asia_font(run, "宋体", 12)

    doc.add_paragraph("")

    sections = [
        (
            "一、项目阶段性进展概述",
            [
                "本学期项目以“校圈”校园信息与社区平台为核心，目标是围绕学生日常使用场景，整合校园通知、课程信息和社区互动功能，形成一个兼具信息聚合与交流属性的小型应用。中期阶段的工作重点主要集中在需求收敛、系统结构搭建、基础接口开发以及前端页面联调等方面。到目前为止，项目已经完成了用户登录、信息中心基础页面、部分通知流展示、后台数据结构设计以及前后端基本通信链路的搭建。",
                "在实际推进过程中，我逐渐认识到，一个课程项目虽然规模不大，但如果涉及多端页面、后端接口、第三方系统接入和数据同步逻辑，整体复杂度并不低。尤其是当最初的功能设想比较丰满时，如果没有及时进行优先级裁剪，就很容易出现“功能写了不少，但真正跑通的闭环不够稳定”的情况。因此，中期阶段对我来说不仅是一个完成代码的过程，也是一个不断修正目标、收缩边界、提高交付可行性的过程。",
            ],
        ),
        (
            "二、AI辅助编程的实际作用与反思",
            [
                "在本项目开发过程中，AI 辅助编程工具起到了非常明显的提效作用。与传统完全依赖手工检索资料、逐段编写代码的方式相比，AI 更像是一名高效率的技术协作者，能够在我提出明确需求后，快速给出接口设计建议、数据库字段草案、前后端页面结构、异常处理思路以及调试路径。这种辅助在项目初期尤其有效，因为在系统骨架尚未完全成形时，开发者往往需要频繁在“想法”与“可执行实现”之间切换，AI 可以显著缩短从方案到原型的距离。",
                "具体来说，AI 在三个方面帮助最大。第一，是快速搭建样板代码。例如信息模块、订阅中心、通知卡片等功能在设计之初需要同时考虑数据结构、接口返回格式以及前端展示字段，AI 可以先给出一套基本可运行的模板，我再结合实际项目结构进行改写，这比完全从零开始要高效得多。第二，是辅助定位问题。在 Canvas 接入过程中，表面上看是“同步成功但没有数据”，实际上问题可能出在登录跳转、统一认证、页面解析或权限校验等多个环节。AI 在分析日志和拆分问题路径方面帮助很大，它会提醒我先验证跳转链路，再验证登录态，最后再看页面结构和数据提取逻辑。第三，是帮助整理技术思路。当需求比较杂乱时，AI 能把问题重新组织成更有层次的实现方案，例如把公共信息源和用户私有信息源分开建模，把“页面抓取”和“官方 API 调用”看作不同策略，而不是混在一起处理。",
                "不过，AI 的作用并不是“替代开发”，而更像“加速试错”和“提供候选解”。实际使用中我也感受到，AI 给出的方案并不总是完全适配本项目的真实环境，尤其是在涉及学校统一认证、第三方页面结构、项目已有代码约束等问题时，AI 很容易建立在一种理想化假设上。如果开发者不做验证，直接照搬，就会把错误快速放大。例如在 Canvas 对接早期，如果只看接口设计，似乎使用 token 是最规范的方案，但对课程项目而言，真实前提是学生不知道 token、学校也不会开放正式对接权限，因此这条技术路径虽然“看起来标准”，却并不符合实际场景。",
                "因此，我对 AI 辅助编程的最大体会是：它非常适合用来加速“分析、搭架子、写初稿、查漏补缺”，但真正决定项目质量的仍然是开发者自己对需求边界、运行环境和风险条件的判断。只有在“先理解项目，再使用 AI”的前提下，AI 才能真正发挥价值。",
            ],
        ),
        (
            "三、项目风险识别与控制",
            [
                "中期开发过程中，我逐渐意识到课程项目最容易被忽视的一部分并不是功能实现本身，而是风险控制。很多问题并不是在编码当天才出现，而是在需求设定阶段就已经埋下了隐患。如果一开始没有识别出来，后续就会不断返工。",
                "本项目目前暴露出的风险主要有三类。第一类是外部依赖风险。无论是微信公众号历史推文抓取、教务系统公告轮询，还是 Canvas 通知同步，本质上都依赖外部平台页面结构、权限机制或登录流程。一旦对方页面改版、增加认证、限制频率，原有方案就可能失效。这类风险的控制思路不是“保证永远稳定”，而是尽量设计替代路径。例如展示层只做聚合和跳转，不复制内容；对公共信息源保留手动补录兜底；对 Canvas 这类个人数据来源采用“用户手动同步”的保守策略，而不是一开始就承诺实时、全自动更新。",
                "第二类是需求膨胀风险。项目一开始的设想包含订阅中心、智能卡片、通知社区讨论联动等多个层次，如果每个模块都追求完整实现，最终很可能导致每个功能都只做到一半。因此我在中期阶段开始更重视“主线闭环”而不是“功能数量”。目前的控制策略是先保证核心链路可演示，例如用户能看到信息流、能完成基本订阅、能理解通知卡片的价值；而对高复杂度功能，例如成熟的第三方系统同步、复杂推荐逻辑等，则放在后续迭代。",
                "第三类是数据与账号安全风险。特别是在处理 Canvas 登录信息时，我意识到课程项目不能简单照搬企业级的“接入第三方账号体系”思路。如果让用户直接提交学号和密码，系统就必须考虑凭证保存、会话有效期、异常退出和最小暴露面等问题。虽然本项目只做课程演示，但依然需要基本的风险意识。因此，当前方案尽量限制抓取结果仅对当前用户可见，并将这一功能明确定位为实验性同步能力，而不是正式对接服务。这种做法虽然保守，但更符合课程项目的实际边界。",
            ],
        ),
        (
            "四、进度管理方法与执行情况",
            [
                "在进度管理方面，我最初采用的是较为粗放的方式，即先把所有功能点列出来，再按模块推进。但实践后发现，这种方式容易陷入“局部任务很多、整体节奏不清晰”的状态。中期之后，我开始更重视阶段划分和任务拆分，把项目进度分为“需求确认—基础架构—核心功能打通—外部能力接入—测试与展示”几个阶段，并为每个阶段设定更具体的可交付目标。",
                "例如在信息模块开发中，我没有把“接入公众号、教务、Canvas”看成同一层任务，而是进一步拆分为：界面先用 mock 数据跑通，后端抽象出统一来源模型，再分别尝试接入真实来源。在这之中，真正影响中期成果展示的不是“是否全部接完”，而是“是否把接入路径想清楚，并且证明系统结构支持后续扩展”。这种拆分方式让我在面对复杂问题时不至于完全停滞，因为即使某个外部系统短期无法打通，其他部分也能继续推进。",
                "同时，我也开始关注里程碑管理的意义。课程项目时间有限，如果没有阶段性的演示目标，很容易在一些细枝末节上投入过多时间。以 Canvas 功能为例，最初如果执着于一步到位完成自动化登录、稳定抓取和完整解析，可能会压缩其他功能的开发空间。中期阶段更合理的做法是先完成可用框架、明确风险、留下调试证据，再根据剩余时间决定是否继续深挖。这种“先完成可展示成果，再追求完善”的思路，对课程项目尤其重要。",
                "当然，从执行效果来看，我的进度管理仍有可以改进的地方。一是前期对外部依赖的复杂度预估不足，导致在部分功能上花费了超出预期的调试时间；二是有些需求在设计时没有立即进行优先级筛选，造成中间出现过返工。之后我需要进一步强化“先评估风险、再排期”的意识，把有限时间优先投向对最终展示价值最高的功能。",
            ],
        ),
        (
            "五、当前问题与后续计划",
            [
                "从当前阶段看，项目已经具备了继续推进的基础，但还存在若干关键问题需要在后续解决。首先，外部系统接入的稳定性仍然不足，特别是涉及统一认证和页面抓取的部分，后续需要在真实登录链路识别、页面结构解析和异常处理方面继续补强。其次，通知卡片的结构化提取目前更多依赖规则化处理，后续可以考虑结合更稳定的关键词抽取和字段识别方式，提升截止时间、适用对象、跳转链接等信息的提取准确率。最后，前端展示虽然已经具备基本功能，但在交互一致性、状态反馈和异常提示上仍有提升空间。",
                "接下来的计划，我会坚持“先稳定主线，再扩展细节”的思路。具体而言，一是继续完善信息模块，让公共通知流、个人通知流和筛选逻辑更加清晰；二是对 Canvas 这类高风险外部接入功能采取分阶段验证策略，优先确认登录链路与页面结构，再决定是否深入自动化同步；三是为项目整理更规范的演示文档和测试说明，确保最终答辩时不仅能展示界面，也能说明系统设计、风险判断和实现权衡。",
            ],
        ),
        (
            "六、中期总结",
            [
                "通过本阶段的开发，我最大的收获并不只是写出了多少代码，而是对“如何在有限时间内完成一个可交付的软件项目”有了更具体的理解。AI 辅助编程确实提升了开发效率，但它真正的价值在于帮助我更快地试错、组织问题和搭建骨架，而不是替代判断。项目风险控制让我意识到，做项目不能只盯着功能点本身，更要识别依赖条件、实现边界和潜在返工成本。进度管理则让我认识到，课程项目最重要的不是追求面面俱到，而是在明确主线目标的前提下，持续产出可验证、可展示、可说明的阶段性成果。",
                "总体来看，这份中期总结不仅是对当前开发工作的回顾，也帮助我进一步明确了后续推进方向。后面的开发中，我会继续保持“需求务实、实现可控、风险透明、节奏清晰”的原则，在现有基础上把项目逐步完善。",
            ],
        ),
    ]

    for title, paragraphs in sections:
        add_heading(doc, title, 1)
        for paragraph in paragraphs:
            add_body(doc, paragraph)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_east_asia_font(run, "宋体", 10)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    run = footer.add_run(" 页")
    set_east_asia_font(run, "宋体", 10)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
